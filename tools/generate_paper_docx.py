#!/usr/bin/env python3
"""
论文 DOCX 生成工具

将 Markdown 论文转为格式规范的 .docx 文件，符合中文本科课程论文排版要求。

依赖：pip install python-docx

用法：
    python generate_paper_docx.py paper.md -o 论文标题.docx
    python generate_paper_docx.py paper.md --style custom_style.json
"""

import argparse
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("[ERROR] 请先安装 python-docx：pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# 默认样式配置
# ---------------------------------------------------------------------------

DEFAULT_STYLE = {
    # 页面设置
    "page": {
        "top_margin": Cm(2.5),
        "bottom_margin": Cm(2.5),
        "left_margin": Cm(2.8),
        "right_margin": Cm(2.8),
    },
    # 论文标题：黑体小三加粗居中，单倍行距，段前段后1行
    "title": {
        "font_name": "黑体",
        "font_size": Pt(15),       # 小三
        "bold": True,
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "space_before": Pt(12),
        "space_after": Pt(12),
        "line_spacing": 1.0,
    },
    # 一级标题：黑体小三加粗左顶格，段前段后0.5行，单倍行距
    "h1": {
        "font_name": "黑体",
        "font_size": Pt(15),       # 小三
        "bold": True,
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(6),
        "space_after": Pt(6),
        "line_spacing": 1.0,
        "page_break_before": False,
    },
    # 二级标题：黑体四号加粗左顶格，段前段后0.5行，单倍行距
    "h2": {
        "font_name": "黑体",
        "font_size": Pt(14),       # 四号
        "bold": True,
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(6),
        "space_after": Pt(6),
        "line_spacing": 1.0,
    },
    # 三级标题：黑体小四不加粗左顶格，段前段后0.5行，单倍行距
    "h3": {
        "font_name": "黑体",
        "font_size": Pt(12),       # 小四
        "bold": False,
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(6),
        "space_after": Pt(6),
        "line_spacing": 1.0,
    },
    # 正文：宋体小四两端对齐，固定行距18磅，首行缩进2字符
    "body": {
        "font_name": "宋体",
        "font_size": Pt(12),       # 小四
        "bold": False,
        "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "line_spacing": Pt(18),
        "first_line_indent": Cm(0.74),
    },
    # 摘要标题：黑体小四加粗
    "abstract_title": {
        "font_name": "黑体",
        "font_size": Pt(12),
        "bold": True,
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "space_after": Pt(0),
    },
    # 关键词标签
    "keywords_label": {
        "font_name": "黑体",
        "font_size": Pt(12),
        "bold": True,
    },
    # 参考文献：宋体小四，固定行距18磅，段前段后0
    "reference": {
        "font_name": "宋体",
        "font_size": Pt(12),       # 小四（模板要求）
        "bold": False,
        "line_spacing": Pt(18),
        "first_line_indent": Cm(0),
    },
    # 图题
    "figure_caption": {
        "font_name": "宋体",
        "font_size": Pt(10.5),
        "bold": False,
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
    },
    # 表题
    "table_caption": {
        "font_name": "宋体",
        "font_size": Pt(10.5),
        "bold": True,
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
    },
}


# ---------------------------------------------------------------------------
# 段落样式应用
# ---------------------------------------------------------------------------

def set_run_font(run, font_name: str, font_size, bold: bool = False):
    """设置 run 的字体和大小"""
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    # 设置中文字体（处理英文默认字体问题）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def apply_paragraph_style(paragraph, style_config: dict):
    """将样式配置应用到段落"""
    pf = paragraph.paragraph_format
    if "alignment" in style_config:
        paragraph.alignment = style_config["alignment"]
    if "space_before" in style_config:
        pf.space_before = style_config["space_before"]
    if "space_after" in style_config:
        pf.space_after = style_config["space_after"]
    if "line_spacing" in style_config:
        pf.line_spacing = style_config["line_spacing"]
    if "first_line_indent" in style_config:
        pf.first_line_indent = style_config["first_line_indent"]
    if "hanging_indent" in style_config:
        pf.first_line_indent = style_config["hanging_indent"]
        # 悬挂缩进 = 首行缩进为负 + 左缩进
        # python-docx 处理方式：左缩进 + 首行缩进负值
        from docx.shared import Cm
        pf.left_indent = Cm(0.74)
        pf.first_line_indent = Cm(-0.74)


def add_styled_paragraph(doc, text: str, style_config: dict, run_config: dict = None):
    """添加一个带样式的段落"""
    p = doc.add_paragraph()
    apply_paragraph_style(p, style_config)
    if run_config is None:
        run_config = style_config
    run = p.add_run(text)
    set_run_font(run, run_config.get("font_name", "宋体"),
                 run_config.get("font_size", Pt(12)),
                 run_config.get("bold", False))
    return p


# ---------------------------------------------------------------------------
# Markdown 解析与转换
# ---------------------------------------------------------------------------

def parse_markdown_to_docx(doc, md_content: str, images_dir: str = None, md_dir: str = None):
    """解析 Markdown 内容并写入 docx 文档"""
    lines = md_content.split("\n")
    i = 0
    in_abstract = False
    in_keywords = False
    in_references = False

    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            i += 1
            continue

        # 一级标题（论文标题）
        if line.startswith("# ") and not line.startswith("## "):
            title_text = line[2:].strip()
            p = add_styled_paragraph(doc, "", DEFAULT_STYLE["title"])
            p.clear()
            run = p.add_run(title_text)
            set_run_font(run, DEFAULT_STYLE["title"]["font_name"],
                        DEFAULT_STYLE["title"]["font_size"], True)
            i += 1
            continue

        # 二级标题（章节标题，如 "## 一、引言"）
        if line.startswith("## "):
            section_title = line[3:].strip()

            # 处理 "摘要" 标题
            if "摘要" in section_title and "abstract" not in section_title.lower():
                p = add_styled_paragraph(doc, section_title, DEFAULT_STYLE["abstract_title"])
                in_abstract = True
            elif section_title.lower().startswith("abstract"):
                p = add_styled_paragraph(doc, section_title, DEFAULT_STYLE["abstract_title"])
            elif "关键词" in section_title or "关键字" in section_title:
                in_keywords = True
                i += 1
                continue
            elif "参考文献" in section_title:
                in_references = True
                p = add_styled_paragraph(doc, section_title, DEFAULT_STYLE["abstract_title"])
                p.paragraph_format.page_break_before = True
            elif section_title.startswith("图") or section_title.startswith("表"):
                # 图题或表题
                if section_title.startswith("图"):
                    style = DEFAULT_STYLE["figure_caption"]
                else:
                    style = DEFAULT_STYLE["table_caption"]
                p = add_styled_paragraph(doc, section_title, style)
            else:
                # 普通一级标题（正文中的一、二、三）
                p = add_styled_paragraph(doc, section_title, DEFAULT_STYLE["h1"])
                in_abstract = False
            i += 1
            continue

        # 三级标题（如 "### 1.1 背景"）
        if line.startswith("### "):
            sub_title = line[4:].strip()
            p = add_styled_paragraph(doc, sub_title, DEFAULT_STYLE["h2"])
            i += 1
            continue

        # 四级标题（如 "#### 2.1.1"）
        if line.startswith("#### "):
            sub_sub = line[5:].strip()
            p = add_styled_paragraph(doc, sub_sub, DEFAULT_STYLE["h3"])
            i += 1
            continue

        # 图片
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_rel_path = img_match.group(2)

            # 多策略路径解析
            resolved_path = None
            # 策略1：images_dir + basename（兼容 --images 参数）
            if images_dir:
                candidate = os.path.join(images_dir, os.path.basename(img_rel_path))
                if os.path.exists(candidate):
                    resolved_path = candidate
            # 策略2：md_dir + 原路径（当图片与 md 在同一项目树下时最可靠）
            if resolved_path is None and md_dir:
                candidate = os.path.join(md_dir, img_rel_path)
                if os.path.exists(candidate):
                    resolved_path = candidate
            # 策略3：原样路径
            if resolved_path is None and os.path.exists(img_rel_path):
                resolved_path = img_rel_path

            if resolved_path and os.path.exists(resolved_path):
                try:
                    # 用 PIL 获取图片实际尺寸以计算合适宽度
                    img_width = Inches(5.5)
                    try:
                        from PIL import Image
                        im = Image.open(resolved_path)
                        w_px, h_px = im.size
                        # 目标最大宽度 5.5in，但高图（表格截图）限制为 4.5in
                        if h_px > w_px * 1.5:
                            img_width = Inches(4.5)
                        elif w_px > h_px * 1.5:
                            img_width = Inches(5.8)
                        im.close()
                    except Exception:
                        pass

                    # 判断是图还是表：alt_text 以"表"开头 → 表题，否则 → 图题
                    is_table = alt_text.startswith("表")

                    # 题注放在图/表上方还是下方
                    # 表题在表上方（加粗居中），图题在图下方（不加粗居中）
                    if is_table:
                        add_styled_paragraph(doc, alt_text, DEFAULT_STYLE["table_caption"])

                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(resolved_path, width=img_width)

                    if not is_table:
                        add_styled_paragraph(doc, alt_text, DEFAULT_STYLE["figure_caption"])
                except Exception as e:
                    print(f"[WARN] 图片插入失败: {resolved_path} — {e}", file=sys.stderr)
                    placeholder = f"[图表: {alt_text}]" if alt_text else "[图表占位]"
                    add_styled_paragraph(doc, placeholder, DEFAULT_STYLE["figure_caption"])
            else:
                print(f"[WARN] 图片未找到: {img_rel_path} (images_dir={images_dir}, md_dir={md_dir})", file=sys.stderr)
                placeholder = f"[图表: {alt_text}]" if alt_text else "[图表占位]"
                add_styled_paragraph(doc, placeholder, DEFAULT_STYLE["figure_caption"])
            i += 1
            continue

        # 表格（简易支持）
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                _insert_table(doc, table_lines)
            continue

        # 关键词行
        if in_keywords and ("关键词" in line or "关键字" in line):
            p = doc.add_paragraph()
            kw_match = re.match(r'\*?\*?关键词\*?\*?\s*[：:]\s*(.*)', line)
            if kw_match:
                kw_text = kw_match.group(1)
                label_run = p.add_run("关键词：")
                set_run_font(label_run, "黑体", Pt(12), True)
                content_run = p.add_run(kw_text)
                set_run_font(content_run, "宋体", Pt(12), False)
            else:
                p.add_run(line)
            in_keywords = False
            i += 1
            continue

        # 参考文献条目
        if in_references and re.match(r'^\[\d+\]', line.strip()):
            ref_text = line.strip()
            p = doc.add_paragraph()
            apply_paragraph_style(p, DEFAULT_STYLE["reference"])
            run = p.add_run(ref_text)
            set_run_font(run, "宋体", Pt(10.5), False)
            i += 1
            continue

        # 普通正文段落
        if in_references and line.strip():
            # 参考文献续行
            ref_text = line.strip()
            p = doc.add_paragraph()
            apply_paragraph_style(p, DEFAULT_STYLE["reference"])
            run = p.add_run(ref_text)
            set_run_font(run, "宋体", Pt(10.5), False)
        else:
            # 移除行内 Markdown 标记
            clean = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', line)
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            clean = clean.strip()

            if clean:
                if in_abstract:
                    style_config = dict(DEFAULT_STYLE["body"])
                    style_config["first_line_indent"] = Cm(0.74)
                    p = add_styled_paragraph(doc, clean, style_config)
                else:
                    p = add_styled_paragraph(doc, clean, DEFAULT_STYLE["body"])

        i += 1


def _insert_table(doc, table_lines: list):
    """将 Markdown 表格行插入到 docx"""
    # 解析表头
    header_cells = [c.strip() for c in table_lines[0].strip("|").split("|")]
    # 跳过分隔线
    data_start = 1
    if len(table_lines) > 1 and re.match(r'^[\|\s\-:]+$', table_lines[1]):
        data_start = 2

    data_rows = []
    for line in table_lines[data_start:]:
        cells = [c.strip() for c in line.strip("|").split("|")]
        data_rows.append(cells)

    cols = len(header_cells)
    rows = 1 + len(data_rows)

    table = doc.add_table(rows=rows, cols=cols, style="Table Grid")
    # 填充表头
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10.5)

    # 填充数据
    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10.5)


# ---------------------------------------------------------------------------
# 主函数
# ---------------------------------------------------------------------------

def generate_docx(md_path: str, output_path: str = None, images_dir: str = None):
    """主入口：将 Markdown 论文转为 DOCX"""
    if not os.path.exists(md_path):
        print(f"[ERROR] 文件不存在: {md_path}", file=sys.stderr)
        sys.exit(1)

    if output_path is None:
        base = os.path.splitext(md_path)[0]
        output_path = f"{base}.docx"

    md_dir = os.path.dirname(os.path.abspath(md_path))

    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    doc = Document()

    # 页面设置
    for section in doc.sections:
        for key, val in DEFAULT_STYLE["page"].items():
            setattr(section, key, val)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    parse_markdown_to_docx(doc, md_content, images_dir, md_dir)

    doc.save(output_path)
    print(f"[OK] DOCX 已生成: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Markdown 论文 → DOCX 生成工具")
    parser.add_argument("input", help="Markdown 论文文件路径")
    parser.add_argument("--output", "-o", help="输出 .docx 路径（默认与输入同名）")
    parser.add_argument("--images", "-i", help="图表图片目录路径")

    args = parser.parse_args()
    generate_docx(args.input, args.output, args.images)


if __name__ == "__main__":
    main()
