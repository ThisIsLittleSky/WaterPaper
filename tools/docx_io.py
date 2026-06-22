#!/usr/bin/env python3
"""Docx I/O for AIGC-Detector Skill.

Commands:
  read <file>                    Print all paragraphs with numbered labels
  replace <file> <idx>            Replace paragraph <idx> (1-based) from stdin
  write <file>                   Write plain text from stdin to new .docx
  analyze <file>                 Extract template formatting metadata
  formatted_write <file>         Write formatted Markdown to .docx (--template for format source)
  insert_figure <file> <idx> <img>  Insert image + caption after paragraph <idx>
"""

import sys
import os
import re
import shutil


def read_docx(file_path: str) -> str:
    """Extract numbered plain text from a .docx file."""
    from docx import Document
    doc = Document(file_path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append(f"[{i + 1}] {text}")
    return "\n\n".join(paragraphs)


def replace_paragraph(file_path: str, index: int, new_text: str, output_path: str = None) -> str:
    """Replace a single paragraph in-place, preserving original formatting.

    Args:
        file_path: Path to source .docx
        index: 1-based paragraph index to replace
        new_text: Replacement text
        output_path: Output path (defaults to {file}_rewritten.docx)
    """
    from docx import Document
    from docx.oxml.ns import qn

    if output_path is None:
        base, ext = os.path.splitext(file_path)
        output_path = f"{base}_rewritten{ext}"

    # Copy to preserve original (skip if input and output are the same file)
    if os.path.abspath(file_path) != os.path.abspath(output_path):
        shutil.copy2(file_path, output_path)

    doc = Document(output_path)
    if index < 1 or index > len(doc.paragraphs):
        print(f"Error: paragraph index {index} out of range (1-{len(doc.paragraphs)})", file=sys.stderr)
        # Clean up failed copy
        if output_path != file_path:
            os.remove(output_path)
        sys.exit(1)

    para = doc.paragraphs[index - 1]

    # Save format from first run BEFORE clearing
    # Merge all w:rPr elements (some .docx files split them across multiple rPr)
    orig_size = None
    orig_name = None
    orig_ea = None
    orig_bold = None
    for r in para.runs:
        if r.text.strip():
            orig_size = r.font.size
            orig_name = r.font.name
            orig_bold = r.font.bold
            for rpr in r._element.findall(qn("w:rPr")):
                rf = rpr.find(qn("w:rFonts"))
                if rf is not None:
                    if orig_ea is None:
                        orig_ea = rf.get(qn("w:eastAsia"))
            break

    style = para.style
    para.clear()
    run = para.add_run(new_text)

    # Restore saved format
    if orig_size is not None:
        run.font.size = orig_size
    if orig_name is not None:
        run.font.name = orig_name
    if orig_bold is not None:
        run.font.bold = orig_bold
    if orig_ea is not None:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = run._element.makeelement(qn("w:rPr"), {})
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:eastAsia"), orig_ea)

    para.style = style

    doc.save(output_path)
    return output_path


def write_docx(file_path: str, text: str) -> None:
    """Write plain text to a new .docx file."""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.name = "宋体"
    doc.save(file_path)


def _parse_pt(s):
    """Parse pt value string for sorting. Returns -1 for 'inherit'."""
    if s == "inherit":
        return -1
    try:
        return float(s.replace("pt", ""))
    except (ValueError, AttributeError):
        return -1


def _get_font_props(run):
    """Extract font properties from a run, handling multiple w:rPr elements.

    Some tools produce runs with multiple w:rPr elements (e.g., one for rFonts,
    another for sz/b). This helper merges all of them.
    """
    from docx.oxml.ns import qn

    size_emu = run.font.size
    bold = run.font.bold
    ea = ascii_f = hAnsi = None

    for rpr in run._element.findall(qn("w:rPr")):
        # Font size from w:sz (half-points → EMU: halfpt * 6350 = EMU)
        if size_emu is None:
            sz = rpr.find(qn("w:sz"))
            if sz is not None:
                val = sz.get(qn("w:val"))
                if val:
                    size_emu = int(val) * 6350
        # Bold from w:b
        if bold is None:
            b_el = rpr.find(qn("w:b"))
            if b_el is not None:
                bold = True
        # Fonts from w:rFonts
        if ea is None or ascii_f is None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                if ea is None:
                    ea = rfonts.get(qn("w:eastAsia"))
                if ascii_f is None:
                    ascii_f = rfonts.get(qn("w:ascii"))
                if hAnsi is None:
                    hAnsi = rfonts.get(qn("w:hAnsi"))

    return size_emu, bold, ea, ascii_f


def analyze_docx(file_path: str) -> str:
    """Extract template formatting metadata from a .docx file.

    Outputs structured text describing page layout, font patterns,
    paragraph formats, and document structure (chapter headings).
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(file_path)
    lines = []

    # --- Page Layout (from XML, more reliable than python-docx properties) ---
    lines.append("=== PAGE LAYOUT ===")
    section = doc.sections[0]
    sectPr = section._sectPr

    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is not None:
        w_twips = int(pgSz.get(qn("w:w"), "0"))
        h_twips = int(pgSz.get(qn("w:h"), "0"))
        w_cm = w_twips / 1440 * 2.54 if w_twips else 21.0
        h_cm = h_twips / 1440 * 2.54 if h_twips else 29.7
        lines.append(f"page_width: {w_cm:.2f} cm")
        lines.append(f"page_height: {h_cm:.2f} cm")

    pgMar = sectPr.find(qn("w:pgMar"))
    if pgMar is not None:
        for attr, label in [
            (qn("w:top"), "margin_top"),
            (qn("w:bottom"), "margin_bottom"),
            (qn("w:left"), "margin_left"),
            (qn("w:right"), "margin_right"),
            (qn("w:gutter"), "gutter"),
        ]:
            val = pgMar.get(attr)
            if val:
                cm = int(val) / 1440 * 2.54
                lines.append(f"{label}: {cm:.2f} cm")

    # --- Font Patterns (group by visual appearance) ---
    lines.append("")
    lines.append("=== FONT PATTERNS ===")
    pattern_map = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        for run in para.runs:
            if run.text.strip():
                size_emu, bold, ea, ascii_f = _get_font_props(run)
                size_val = f"{size_emu / 12700:.1f}pt" if size_emu else "inherit"
                align = str(para.paragraph_format.alignment) if para.paragraph_format.alignment else "left"

                key = (size_val, bold, align, ea, ascii_f)
                if key not in pattern_map:
                    pattern_map[key] = []
                sample = text[:60]
                if not any(sample in s for s in pattern_map[key]):
                    if len(pattern_map[key]) < 3:
                        pattern_map[key].append(sample)
                break

    for key, samples in sorted(pattern_map.items(), key=lambda x: -_parse_pt(x[0][0])):
        size_val, bold, align, ea, ascii_f = key
        lines.append(f"Pattern: size={size_val}, bold={bold}, align={align}, "
                      f"eastAsia={ea}, ascii={ascii_f}")
        for s in samples:
            lines.append(f"  sample: \"{s}\"")

    # --- Paragraph Formats ---
    lines.append("")
    lines.append("=== PARAGRAPH FORMATS ===")
    for para in doc.paragraphs[:30]:
        text = para.text.strip()
        if not text or len(text) < 20:
            continue
        pf = para.paragraph_format
        ls = pf.line_spacing
        fi = pf.first_line_indent
        sa = pf.space_after
        sb = pf.space_before
        if any(v is not None for v in [ls, fi, sa, sb]):
            lines.append(f"  text=\"{text[:50]}\" line_spacing={ls} "
                          f"first_indent={fi} space_before={sb} space_after={sa}")

    # --- Document Structure Detection ---
    lines.append("")
    lines.append("=== DOCUMENT STRUCTURE ===")
    ch_patterns = [
        (r"^第[一二三四五六七八九十]+章\s", "CN-chapter"),
        (r"^第[一二三四五六七八九十]+节\s", "CN-section"),
        (r"^\d+(\.\d+)*\s+\S", "NUM-section"),
        (r"^Chapter\s+\d+", "EN-chapter"),
        (r"^(Abstract|摘要|关键词|Keywords|目录|参考文献|致谢|Acknowledgments"
          r"|前言|绪论|总结|结论|附录|引言)", "SPECIAL"),
    ]
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pattern, label in ch_patterns:
            if re.match(pattern, text):
                for run in para.runs:
                    if run.text.strip():
                        size_emu, bold, ea, ascii_f = _get_font_props(run)
                        size_pt = f"{size_emu / 12700:.1f}pt" if size_emu else "inherit"
                        lines.append(f"  [{i}] {label}: \"{text[:60]}\" "
                                      f"(size={size_pt}, bold={bold})")
                        break
                break

    return "\n".join(lines)


def _extract_template_formats(tdoc):
    """Extract formatting rules from a template document.

    Returns a dict with keys: title, heading1, heading2, heading3, body.
    Each value is a dict of format properties.
    """
    from docx.oxml.ns import qn

    # Collect all paragraph format signatures grouped by font size
    size_groups = {}
    for para in tdoc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        for run in para.runs:
            if run.text.strip():
                size_emu, bold, ea, ascii_f = _get_font_props(run)
                size_pt = size_emu / 12700 if size_emu else 12.0
                align = para.paragraph_format.alignment
                pf = para.paragraph_format

                key = round(size_pt, 1)
                if key not in size_groups:
                    size_groups[key] = {
                        "size_pt": size_pt,
                        "bold": bold,
                        "alignment": align,
                        "eastAsia": ea,
                        "ascii": ascii_f,
                        "line_spacing": pf.line_spacing,
                        "first_line_indent": pf.first_line_indent,
                        "space_before": pf.space_before,
                        "space_after": pf.space_after,
                        "count": 0,
                        "sample": text[:40],
                    }
                size_groups[key]["count"] += 1
                break

    # Sort by font size descending to identify heading levels
    sorted_sizes = sorted(size_groups.items(), key=lambda x: -x[0])

    fmt = {}
    for i, (size, props) in enumerate(sorted_sizes):
        if i == 0 and props["count"] <= 3 and props["size_pt"] >= 16:
            fmt["title"] = props
            continue
        # Assign remaining levels by position (skip body once reached)
        if "title" in fmt and i == 1 and "heading1" not in fmt:
            fmt["heading1"] = props
        elif "heading1" not in fmt:
            fmt["heading1"] = props
        elif "heading2" not in fmt:
            fmt["heading2"] = props
        elif "heading3" not in fmt:
            fmt["heading3"] = props
        else:
            fmt["body"] = props
            break

    # Ensure body exists (use most common small size)
    if "body" not in fmt and sorted_sizes:
        # Find the smallest size group with the most paragraphs
        body_candidates = sorted(
            [(s, p) for s, p in sorted_sizes if p["size_pt"] <= 14],
            key=lambda x: -x[1]["count"]
        )
        if body_candidates:
            fmt["body"] = body_candidates[0][1]
        else:
            # Fallback defaults
            fmt["body"] = {
                "size_pt": 12.0, "bold": None, "alignment": 3,
                "eastAsia": "宋体", "ascii": "Times New Roman",
                "line_spacing": 1.5, "first_line_indent": None,
                "space_before": None, "space_after": None,
            }

    # Defaults for missing levels
    if "heading1" not in fmt:
        fmt["heading1"] = {**fmt["body"], "size_pt": 16.0, "bold": True, "alignment": 1}
    if "heading2" not in fmt:
        fmt["heading2"] = {**fmt["body"], "size_pt": 14.0, "bold": True}
    if "heading3" not in fmt:
        fmt["heading3"] = {**fmt["body"], "size_pt": 12.0, "bold": True}

    return fmt


def _copy_page_layout(doc, template_path: str):
    """Copy page layout from template to target document."""
    from docx import Document
    from docx.oxml.ns import qn

    tdoc = Document(template_path)
    src_sectPr = tdoc.sections[0]._sectPr
    dst_section = doc.sections[0]
    dst_sectPr = dst_section._sectPr

    # Copy pgSz
    pgSz = src_sectPr.find(qn("w:pgSz"))
    if pgSz is not None:
        existing = dst_sectPr.find(qn("w:pgSz"))
        if existing is not None:
            dst_sectPr.remove(existing)
        dst_sectPr.append(pgSz)

    # Copy pgMar
    pgMar = src_sectPr.find(qn("w:pgMar"))
    if pgMar is not None:
        existing = dst_sectPr.find(qn("w:pgMar"))
        if existing is not None:
            dst_sectPr.remove(existing)
        dst_sectPr.append(pgMar)


def _apply_format(paragraph, props, fmt):
    """Apply formatting properties to a paragraph and its runs."""
    from docx.shared import Pt, Emu, Cm
    from docx.oxml.ns import qn

    if not props:
        return

    # Paragraph alignment
    if props.get("alignment") is not None:
        paragraph.paragraph_format.alignment = props["alignment"]

    # Line spacing
    ls = props.get("line_spacing")
    if ls is not None:
        paragraph.paragraph_format.line_spacing = ls

    # First line indent (2 chars ≈ 24pt for 12pt font)
    if props.get("first_line_indent") is not None:
        paragraph.paragraph_format.first_line_indent = props["first_line_indent"]

    # Space before/after
    sb = props.get("space_before")
    sa = props.get("space_after")
    if sb is not None:
        paragraph.paragraph_format.space_before = sb
    if sa is not None:
        paragraph.paragraph_format.space_after = sa

    # Apply font properties to each run
    for run in paragraph.runs:
        size_pt = props.get("size_pt")
        if size_pt:
            run.font.size = Pt(size_pt)
        if props.get("bold"):
            run.font.bold = True
        elif props.get("bold") is False:
            run.font.bold = False

        # East Asian font via XML
        ea = props.get("eastAsia")
        ascii_f = props.get("ascii")
        if ea or ascii_f:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is None:
                rpr = run._element.makeelement(qn("w:rPr"), {})
                run._element.insert(0, rpr)
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.insert(0, rfonts)
            if ea:
                rfonts.set(qn("w:eastAsia"), ea)
            if ascii_f:
                rfonts.set(qn("w:ascii"), ascii_f)
                rfonts.set(qn("w:hAnsi"), ascii_f)


def _add_markdown_runs(paragraph, text, props, fmt):
    """Parse **bold** and *italic* markers in text and add runs accordingly."""
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            # Bold + Italic
            run = paragraph.add_run(part[3:-3])
            run.font.bold = True
            run.font.italic = True
        elif part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            # Plain text
            paragraph.add_run(part)
    # Apply base format to all runs
    if props:
        for run in paragraph.runs:
            size_pt = props.get("size_pt")
            if size_pt:
                run.font.size = Pt(size_pt)
            ea = props.get("eastAsia")
            ascii_f = props.get("ascii")
            if ea or ascii_f:
                rpr = run._element.find(qn("w:rPr"))
                if rpr is None:
                    rpr = run._element.makeelement(qn("w:rPr"), {})
                    run._element.insert(0, rpr)
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = rpr.makeelement(qn("w:rFonts"), {})
                    rpr.insert(0, rfonts)
                if ea:
                    rfonts.set(qn("w:eastAsia"), ea)
                if ascii_f:
                    rfonts.set(qn("w:ascii"), ascii_f)
                    rfonts.set(qn("w:hAnsi"), ascii_f)


def formatted_write_docx(file_path: str, text: str, template_path: str = None):
    """Write formatted Markdown text to a .docx file.

    Parses format markers:
      <!-- thesis-title --> ... <!-- /thesis-title -->
      <!-- chapter-N --> ... <!-- /chapter-N -->
      <!-- figure-placeholder --> ... <!-- /figure-placeholder -->
      <!-- table-placeholder --> ... <!-- /table-placeholder -->

    Markdown headings (# ## ###) map to heading1/heading2/heading3 formatting.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Extract template formatting if provided
    fmt = {}
    if template_path and os.path.exists(template_path):
        from docx import Document as TDoc
        tdoc = TDoc(template_path)
        fmt = _extract_template_formats(tdoc)
        _copy_page_layout(doc, template_path)

    # Parse and write content
    lines = text.strip().split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Skip HTML format markers (they're structural, not content)
        if re.match(r"<!--\s*/?(/\s*)?(thesis-title|chapter-\d+|"
                     r"figure-placeholder|table-placeholder)\s*/?-->", stripped):
            continue

        # Detect Markdown headings
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            _add_markdown_runs(p, stripped[4:], fmt.get("heading3"), fmt)
            _apply_format(p, fmt.get("heading3"), fmt)
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            _add_markdown_runs(p, stripped[3:], fmt.get("heading2"), fmt)
            _apply_format(p, fmt.get("heading2"), fmt)
        elif stripped.startswith("# "):
            p = doc.add_paragraph()
            _add_markdown_runs(p, stripped[2:], fmt.get("heading1"), fmt)
            _apply_format(p, fmt.get("heading1"), fmt)
        else:
            # Body text (supports **bold** and *italic*)
            p = doc.add_paragraph()
            _add_markdown_runs(p, stripped, fmt.get("body"), fmt)
            _apply_format(p, fmt.get("body"), fmt)

    # Ensure output directory exists
    out_dir = os.path.dirname(file_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc.save(file_path)


def insert_figure(file_path: str, paragraph_index: int, image_path: str,
                  caption: str = "", output_path: str = None) -> str:
    """Insert an image + caption paragraph after the specified paragraph index."""
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document(file_path)
    paragraphs = doc.paragraphs

    if paragraph_index < 1 or paragraph_index > len(paragraphs):
        print(f"Error: paragraph index {paragraph_index} out of range (1-{len(paragraphs)})",
              file=sys.stderr)
        sys.exit(1)

    target_para = paragraphs[paragraph_index - 1]

    # Calculate image width from page layout (default 14cm if unknown)
    image_width = Cm(14)
    try:
        section = doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        if page_width and left_margin and right_margin:
            image_width = page_width - left_margin - right_margin
    except Exception:
        pass

    # Create image paragraph after target
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(image_path, width=image_width)

    # Move image paragraph to after target
    target_element = target_para._element
    img_element = img_para._element
    target_element.addnext(img_element)

    # Create caption paragraph if provided
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.font.size = Pt(10.5)
        cap_run.font.name = "宋体"
        cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

        # Move caption to after image
        img_element.addnext(cap_para._element)

    # Save
    out = output_path or file_path
    out_dir = os.path.dirname(out)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(out)

    print(out)
    return out


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 docx_io.py <read|replace|write|analyze|formatted_write|insert_figure> "
              "<file_path> [args...]", file=sys.stderr)
        sys.exit(1)

    command = sys.argv[1]
    file_path = sys.argv[2]

    if command == "read":
        if not os.path.exists(file_path):
            print(f"Error: file not found: {file_path}", file=sys.stderr)
            sys.exit(1)
        text = read_docx(file_path)
        print(text)
    elif command == "replace":
        if len(sys.argv) < 4:
            print("Usage: python3 docx_io.py replace <file_path> <paragraph_index> "
                  "[--output <path>]", file=sys.stderr)
            print("  Paragraph text is read from stdin.", file=sys.stderr)
            sys.exit(1)
        index = int(sys.argv[3])
        out_path = None
        if "--output" in sys.argv:
            oidx = sys.argv.index("--output")
            if oidx + 1 < len(sys.argv):
                out_path = sys.argv[oidx + 1]
        new_text = sys.stdin.read().strip()
        output = replace_paragraph(file_path, index, new_text, output_path=out_path)
        print(output, file=sys.stderr)
    elif command == "write":
        text = sys.stdin.read()
        write_docx(file_path, text)
        print(f"Written to: {file_path}", file=sys.stderr)
    elif command == "analyze":
        if not os.path.exists(file_path):
            print(f"Error: file not found: {file_path}", file=sys.stderr)
            sys.exit(1)
        metadata = analyze_docx(file_path)
        print(metadata)
    elif command == "formatted_write":
        template_path = None
        if "--template" in sys.argv:
            tidx = sys.argv.index("--template")
            if tidx + 1 < len(sys.argv):
                template_path = sys.argv[tidx + 1]
        text = sys.stdin.read()
        formatted_write_docx(file_path, text, template_path)
        print(f"Written to: {file_path}", file=sys.stderr)
    elif command == "insert_figure":
        if len(sys.argv) < 5:
            print("Usage: python3 docx_io.py insert_figure <file_path> <paragraph_index> "
                  "<image_path> [--caption <text>] [--output <path>]", file=sys.stderr)
            sys.exit(1)
        index = int(sys.argv[3])
        image = sys.argv[4]
        caption = ""
        out_path = None
        if "--caption" in sys.argv:
            cidx = sys.argv.index("--caption")
            if cidx + 1 < len(sys.argv):
                caption = sys.argv[cidx + 1]
        if "--output" in sys.argv:
            oidx = sys.argv.index("--output")
            if oidx + 1 < len(sys.argv):
                out_path = sys.argv[oidx + 1]
        if not os.path.exists(image):
            print(f"Error: image not found: {image}", file=sys.stderr)
            sys.exit(1)
        output = insert_figure(file_path, index, image, caption, output_path=out_path)
        print(f"Figure inserted: {output}", file=sys.stderr)
    else:
        print(f"Error: unknown command '{command}'. "
              f"Use 'read', 'replace', 'write', 'analyze', or 'formatted_write'.",
              file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
