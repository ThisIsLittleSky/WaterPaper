#!/usr/bin/env python3
"""
论文模板格式分析工具

分析用户提供的 .docx / .doc 模板文件，提取其排版样式，输出结构化 JSON 供后续写作和 DOCX 生成使用。

支持格式：.docx（原生）、.doc（自动转换为 .docx 后分析）

支持识别：
- 论文标题
- 摘要标题 / Abstract 标题
- 关键词行
- 一级标题（一、二、三...）
- 二级标题（（一）（二）/ 1.1 1.2...）
- 中文正文 / 英文正文
- 图题 / 表题
- 参考文献条目
- 致谢

用法：
    python analyze_template.py template.docx
    python analyze_template.py template.docx --json-out style_profile.json
    python analyze_template.py template.docx --text-out template_text.txt
    python analyze_template.py template.docx --json-out style_profile.json --text-out template_text.txt
    python analyze_template.py template.doc
    python analyze_template.py template.doc --json-out style_profile.json
"""

import json
import os
import re
import subprocess
import sys
import tempfile
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("[ERROR] 请先安装 python-docx：pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# .doc → .docx 转换
# ---------------------------------------------------------------------------

def _convert_via_win32com(doc_path: Path, output_dir: Path) -> Optional[Path]:
    """通过 Microsoft Word COM 自动化将 .doc 转换为 .docx"""
    try:
        import win32com.client
    except ImportError:
        return None

    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
    except Exception:
        return None

    try:
        docx_path = output_dir / f"{doc_path.stem}.docx"
        abs_input = str(doc_path.resolve())
        abs_output = str(docx_path.resolve())

        doc = word.Documents.Open(abs_input, ReadOnly=True)
        doc.SaveAs2(abs_output, FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
        doc.Close()
        return docx_path
    except Exception:
        return None
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def _convert_via_libreoffice(doc_path: Path, output_dir: Path) -> Optional[Path]:
    """通过 LibreOffice 将 .doc 转换为 .docx"""
    candidates = [
        "libreoffice",
        "soffice",
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
    ]

    lo_path = None
    for c in candidates:
        try:
            result = subprocess.run([c, "--version"], capture_output=True,
                                    creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0)
            if result.returncode == 0:
                lo_path = c
                break
        except Exception:
            continue

    if lo_path is None:
        return None

    try:
        abs_output = str(output_dir.resolve())
        abs_input = str(doc_path.resolve())
        subprocess.run(
            [lo_path, "--headless", "--convert-to", "docx",
             "--outdir", abs_output, abs_input],
            capture_output=True, timeout=60,
            creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0,
        )
        docx_path = output_dir / f"{doc_path.stem}.docx"
        if docx_path.exists():
            return docx_path
    except Exception:
        pass
    return None


def convert_doc_to_docx(doc_path: Path) -> Path:
    """
    将 .doc 文件转换为 .docx，返回 .docx 路径。
    优先使用 Word COM 自动化，回退到 LibreOffice。
    """
    suffix = doc_path.suffix.lower()
    if suffix == ".docx":
        return doc_path

    if suffix != ".doc":
        raise ValueError(f"不支持的文件格式: {suffix}，仅支持 .docx / .doc")

    tmpdir = Path(tempfile.mkdtemp(prefix="paper_template_"))

    # 优先尝试 Word COM 自动化（Windows）
    docx_path = _convert_via_win32com(doc_path, tmpdir)
    if docx_path:
        print(f"[OK] 通过 Microsoft Word 将 .doc 转换为 .docx")
        return docx_path

    # 回退到 LibreOffice
    docx_path = _convert_via_libreoffice(doc_path, tmpdir)
    if docx_path:
        print(f"[OK] 通过 LibreOffice 将 .doc 转换为 .docx")
        return docx_path

    # 均不可用
    raise RuntimeError(
        "无法转换 .doc 文件。请安装以下任一软件：\n"
        "  1. Microsoft Word（Windows 上自动检测）\n"
        "  2. LibreOffice（https://www.libreoffice.org/download/）\n"
        "或者手动将 .doc 另存为 .docx 后重新运行。"
    )


# ---------------------------------------------------------------------------
# 标题识别规则（课程论文用"一、二、三"，非毕业论文"第X章"）
# ---------------------------------------------------------------------------

HEADING_PATTERNS = {
    "heading1": re.compile(r"^[一二三四五六七八九十]、"),   # 一、引言
    "heading2": re.compile(r"^（[一二三四五六七八九十]）"),  # （一）小节
    "heading2_alt": re.compile(r"^\d+\.\d+(?!\.)"),          # 1.1 小节
    "heading3": re.compile(r"^\d+\.\d+\.\d+"),                # 1.1.1
}


def alignment_name(value: Optional[int]) -> str:
    mapping = {
        None: "left",
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
    }
    return mapping.get(value, "left")


def east_asia_font(run) -> Optional[str]:
    """获取东亚字体设置"""
    if run._element.rPr is None or run._element.rPr.rFonts is None:
        return None
    return run._element.rPr.rFonts.get(qn("w:eastAsia"))


def paragraph_signature(para) -> dict:
    """提取段落样式指纹"""
    run_fonts = []
    run_sizes = []
    bold_values = []

    for run in para.runs:
        if not run.text.strip():
            continue
        run_fonts.append((east_asia_font(run) or run.font.name or "").strip())
        if run.font.size:
            run_sizes.append(round(run.font.size.pt, 1))
        if run.bold is not None:
            bold_values.append(bool(run.bold))

    pf = para.paragraph_format
    first_line_indent = pf.first_line_indent

    # 检查是否分页
    page_break_before = False
    p_pr = para._p.pPr
    if p_pr is not None and p_pr.pageBreakBefore is not None:
        page_break_before = True

    return {
        "font": most_common(run_fonts),
        "size_pt": most_common(run_sizes),
        "bold": most_common(bold_values, default=False),
        "alignment": alignment_name(para.alignment),
        "first_line_indent_pt": round(first_line_indent.pt, 1) if first_line_indent else 0.0,
        "line_spacing": round(float(pf.line_spacing), 2) if isinstance(pf.line_spacing, (int, float)) else None,
        "space_before_pt": round(pf.space_before.pt, 1) if pf.space_before else 0.0,
        "space_after_pt": round(pf.space_after.pt, 1) if pf.space_after else 0.0,
        "page_break_before": page_break_before,
    }


def most_common(values: list, default: Any = None) -> Any:
    filtered = [v for v in values if v not in ("", None)]
    if not filtered:
        return default
    return Counter(filtered).most_common(1)[0][0]


def classify_paragraph(text: str, in_reference_section: bool, in_ack_section: bool) -> Optional[str]:
    """根据文本内容分类段落类型"""
    normalized = text.strip()
    if not normalized:
        return None

    # 摘要
    if normalized == "摘要":
        return "abstract_heading"
    if normalized.lower() == "abstract":
        return "abstract_heading_en"

    # 关键词
    if normalized.startswith("关键词") or normalized.startswith("关键字"):
        return "keywords"
    if normalized.startswith("Keywords") or normalized.startswith("Key words"):
        return "keywords_en"

    # 参考文献
    if normalized == "参考文献":
        return "references_heading"

    # 致谢
    if normalized == "致谢" or normalized == "谢辞":
        return "acknowledgement_heading"

    # 在参考文献区域内
    if in_reference_section and re.match(r"^\[\d+\]", normalized):
        return "reference_entry"

    # 在致谢区域内
    if in_ack_section:
        return "acknowledgement_body"

    # 图题/表题
    if normalized.startswith("图") and re.match(r"^图\s*\d+", normalized):
        return "figure_caption"
    if normalized.startswith("表") and re.match(r"^表\s*\d+", normalized):
        return "table_caption"

    # 标题层级
    if HEADING_PATTERNS["heading3"].match(normalized):
        return "heading3"
    if HEADING_PATTERNS["heading2_alt"].match(normalized):
        return "heading2"
    if HEADING_PATTERNS["heading2"].match(normalized):
        return "heading2"
    if HEADING_PATTERNS["heading1"].match(normalized):
        return "heading1"

    # 英文段落（连续英文字母且无中文）
    if re.search(r"[A-Za-z]{10,}", normalized) and not re.search(r"[一-鿿]", normalized):
        return "body_en"

    # 默认中文正文
    return "body_cn"


def aggregate_styles(doc: Document) -> dict:
    """遍历文档，聚合各类型段落的样式"""
    grouped = defaultdict(list)
    in_reference_section = False
    in_ack_section = False

    for para in doc.paragraphs:
        text = para.text.strip()

        # 跟踪参考文献区域
        if text == "参考文献":
            in_reference_section = True
            in_ack_section = False
        elif text in ("致谢", "谢辞"):
            in_ack_section = True
            in_reference_section = False
        elif HEADING_PATTERNS["heading1"].match(text):
            # 进入新章节，退出参考文献/致谢区域
            in_reference_section = False
            in_ack_section = False

        category = classify_paragraph(text, in_reference_section, in_ack_section)
        if category is None:
            continue
        grouped[category].append(paragraph_signature(para))

    # 聚合每组样式
    result = {}
    for cat, sigs in grouped.items():
        result[cat] = {
            "font": most_common([s["font"] for s in sigs]),
            "size_pt": most_common([s["size_pt"] for s in sigs]),
            "bold": most_common([s["bold"] for s in sigs], default=False),
            "alignment": most_common([s["alignment"] for s in sigs], default="left"),
            "first_line_indent_pt": most_common([s["first_line_indent_pt"] for s in sigs], default=0.0),
            "line_spacing": most_common([s["line_spacing"] for s in sigs]),
            "space_before_pt": most_common([s["space_before_pt"] for s in sigs], default=0.0),
            "space_after_pt": most_common([s["space_after_pt"] for s in sigs], default=0.0),
            "page_break_before": most_common([s["page_break_before"] for s in sigs], default=False),
            "sample_count": len(sigs),
        }

    # 提取页面设置
    page_style = extract_page_style(doc)
    if page_style:
        result["_page"] = page_style

    return result


def extract_full_text(doc: Document) -> str:
    """导出模板全文，每条段落标注分类标签，供 LLM 分析格式补充信息"""
    lines = []
    in_reference_section = False
    in_ack_section = False

    category_labels = {
        "abstract_heading": "摘要标题",
        "abstract_heading_en": "Abstract标题",
        "keywords": "关键词行",
        "keywords_en": "Keywords行",
        "references_heading": "参考文献标题",
        "reference_entry": "参考文献条目",
        "acknowledgement_heading": "致谢标题",
        "acknowledgement_body": "致谢正文",
        "heading1": "一级标题",
        "heading2": "二级标题",
        "heading3": "三级标题",
        "figure_caption": "图题",
        "table_caption": "表题",
        "body_en": "英文正文",
        "body_cn": "中文正文",
    }

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            # 空段落也记录，帮助 LLM 理解段落间距
            lines.append(f"[空行]")
            continue

        # 跟踪参考文献/致谢区域
        if text == "参考文献":
            in_reference_section = True
            in_ack_section = False
        elif text in ("致谢", "谢辞"):
            in_ack_section = True
            in_reference_section = False
        elif HEADING_PATTERNS["heading1"].match(text):
            in_reference_section = False
            in_ack_section = False

        category = classify_paragraph(text, in_reference_section, in_ack_section)
        label = category_labels.get(category, "未分类") if category else "未分类"
        sig = paragraph_signature(para)

        # 格式： [分类标签] 文字内容  | 字体=X 字号=Ypt 加粗=Z 对齐=W
        meta = f"字体={sig['font'] or '?'} 字号={sig['size_pt']}pt 加粗={'是' if sig['bold'] else '否'} 对齐={sig['alignment']}"
        lines.append(f"[{label}] {text}  |  {meta}")

    return "\n".join(lines)


def extract_page_style(doc: Document) -> dict:
    """提取页面边距设置"""
    try:
        section = doc.sections[0]
        return {
            "top_cm": round(section.top_margin.cm, 2),
            "bottom_cm": round(section.bottom_margin.cm, 2),
            "left_cm": round(section.left_margin.cm, 2),
            "right_cm": round(section.right_margin.cm, 2),
        }
    except Exception:
        return {}


def detect_paper_title(doc: Document) -> Optional[str]:
    """尝试检测论文标题（通常是文档开头居中加粗大字号文本）"""
    for para in list(doc.paragraphs)[:5]:  # 只查前 5 段
        text = para.text.strip()
        if not text or len(text) < 5 or len(text) > 60:
            continue
        sig = paragraph_signature(para)
        if (sig.get("bold") and
            sig.get("alignment") == "center" and
            sig.get("size_pt") and sig["size_pt"] >= 15):
            return text
    return None


def compare_with_defaults(template_styles: dict) -> list:
    """与默认格式对比，生成冲突/差异清单"""
    defaults = {
        "heading1": {"font": "黑体", "size_pt": 15, "bold": True},
        "heading2": {"font": "黑体", "size_pt": 14, "bold": True},
        "body_cn": {"font": "宋体", "size_pt": 12, "bold": False},
        "abstract_heading": {"font": "黑体", "size_pt": 12, "bold": True},
        "keywords": {"font": "黑体", "size_pt": 12, "bold": True},
        "reference_entry": {"font": "宋体", "size_pt": 10.5, "bold": False},
    }

    conflicts = []
    for key, default in defaults.items():
        if key in template_styles:
            t = template_styles[key]
            diffs = []
            if t.get("font") != default.get("font"):
                diffs.append(f"字体: {t['font']} (默认 {default['font']})")
            if t.get("size_pt") != default.get("size_pt"):
                diffs.append(f"字号: {t['size_pt']}pt (默认 {default['size_pt']}pt)")
            if t.get("bold") != default.get("bold"):
                diffs.append(f"加粗: {t['bold']} (默认 {default['bold']})")
            if diffs:
                conflicts.append({"element": key, "differences": diffs})
    return conflicts


# ---------------------------------------------------------------------------
# 输出
# ---------------------------------------------------------------------------

def fmt_summary(styles: dict) -> str:
    """格式化为可读摘要"""
    labels = {
        "heading1": "一级标题（一、）",
        "heading2": "二级标题（1.1）",
        "heading3": "三级标题",
        "body_cn": "中文正文",
        "body_en": "英文正文",
        "abstract_heading": "摘要标题",
        "abstract_heading_en": "Abstract 标题",
        "keywords": "关键词",
        "keywords_en": "Keywords",
        "figure_caption": "图题",
        "table_caption": "表题",
        "reference_entry": "参考文献条目",
        "acknowledgement_body": "致谢正文",
    }
    lines = []
    for key, label in labels.items():
        if key not in styles:
            continue
        s = styles[key]
        lines.append(
            f"  {label:14s} | {s['font'] or '?'} | {s['size_pt']}pt | "
            f"{'加粗' if s['bold'] else '常规'} | "
            f"{'居中' if s['alignment'] == 'center' else '左对齐'} | "
            f"首行缩进{s['first_line_indent_pt']}pt | "
            f"行距{s['line_spacing']} | "
            f"段前{s['space_before_pt']}pt 段后{s['space_after_pt']}pt"
        )
    return "\n".join(lines)


def main():
    if len(sys.argv) < 2:
        print("Usage: python analyze_template.py <template.docx|template.doc> [--json-out path] [--text-out path]")
        return 1

    provided_path = Path(sys.argv[1])
    if not provided_path.exists():
        print(f"[ERROR] 文件不存在: {provided_path}", file=sys.stderr)
        return 1

    # 解析可选参数
    json_out_path = None
    text_out_path = None
    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == "--json-out" and i + 1 < len(sys.argv):
            json_out_path = Path(sys.argv[i + 1])
            i += 2
        elif sys.argv[i] == "--text-out" and i + 1 < len(sys.argv):
            text_out_path = Path(sys.argv[i + 1])
            i += 2
        else:
            i += 1

    # .doc → .docx 自动转换
    docx_path = convert_doc_to_docx(provided_path)

    doc = Document(str(docx_path))
    styles = aggregate_styles(doc)

    # 检测论文标题
    title = detect_paper_title(doc)

    # 构建输出
    output = {
        "source_file": str(docx_path),
        "paper_title": title,
        "styles": styles,
        "conflicts_with_default": compare_with_defaults(styles),
    }

    # JSON 输出
    if json_out_path:
        json_out_path.parent.mkdir(parents=True, exist_ok=True)
        json_out_path.write_text(json.dumps(output, ensure_ascii=False, indent=2),
                                 encoding="utf-8")
        print(f"[OK] 样式配置已保存到: {json_out_path}")

    # 全文导出（供 LLM 分析格式补充信息）
    if text_out_path:
        full_text = extract_full_text(doc)
        text_out_path.parent.mkdir(parents=True, exist_ok=True)
        text_out_path.write_text(full_text, encoding="utf-8")
        print(f"[OK] 模板全文已导出到: {text_out_path}")

    # 终端摘要
    print(f"\n{'='*60}")
    print(f"模板文件: {docx_path}")
    print(f"检测到论文标题: {title or '(未检测到)'}")
    print(f"{'='*60}")
    print(fmt_summary(styles))
    print(f"{'='*60}")

    if output["conflicts_with_default"]:
        print("\n[WARN] 与默认格式的差异:")
        for c in output["conflicts_with_default"]:
            print(f"  {c['element']}:")
            for d in c["differences"]:
                print(f"    - {d}")
    else:
        print("\n[OK] 模板格式与默认格式一致")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
